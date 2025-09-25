"""
PDF conversion utilities for ToolHub.
Handles conversion between PDF and various office formats (Word, Excel, PowerPoint).
"""

import os
import tempfile
from typing import Optional, Dict, Any
import logging

# PDF Libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Office Document Libraries
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from pptx import Presentation
    from pptx.util import Inches as PptxInches
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False


class PDFConverter:
    """Handles conversion between PDF and office formats."""
    
    def __init__(self):
        """Initialize PDF converter with available libraries."""
        self.available_conversions = {
            'pdf_to_word': PYMUPDF_AVAILABLE and DOCX_AVAILABLE,
            'pdf_to_excel': PDFPLUMBER_AVAILABLE and OPENPYXL_AVAILABLE,
            'pdf_to_powerpoint': PYMUPDF_AVAILABLE and PPTX_AVAILABLE,
            'word_to_pdf': DOCX_AVAILABLE and REPORTLAB_AVAILABLE,
            'excel_to_pdf': OPENPYXL_AVAILABLE and REPORTLAB_AVAILABLE,
            'powerpoint_to_pdf': PPTX_AVAILABLE and REPORTLAB_AVAILABLE,
            'text_extraction': PYMUPDF_AVAILABLE or PDFPLUMBER_AVAILABLE
        }
    
    def pdf_to_word(self, pdf_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert PDF to Word document.
        
        Args:
            pdf_path: Path to input PDF
            output_path: Output path (temp file if None)
            
        Returns:
            Path to created Word document
        """
        if not self.available_conversions['pdf_to_word']:
            raise RuntimeError("PDF to Word conversion not available. Install PyMuPDF and python-docx.")
        
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        if output_path is None:
            output_fd, output_path = tempfile.mkstemp(suffix='.docx', prefix='converted_')
            os.close(output_fd)
        
        return self._pdf_to_word_pymupdf(pdf_path, output_path)
    
    def _pdf_to_word_pymupdf(self, pdf_path: str, output_path: str) -> str:
        """Convert PDF to Word using PyMuPDF."""
        doc = fitz.open(pdf_path)
        word_doc = Document()
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text
            text = page.get_text()
            
            if text.strip():
                # Add page break except for first page
                if page_num > 0:
                    word_doc.add_page_break()
                
                # Add text content
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        word_doc.add_paragraph(para.strip())
            
            # Extract images
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                        
                        # Save image temporarily
                        temp_img_fd, temp_img_path = tempfile.mkstemp(suffix='.png')
                        os.close(temp_img_fd)
                        
                        with open(temp_img_path, 'wb') as img_file:
                            img_file.write(img_data)
                        
                        # Add image to Word document
                        try:
                            word_doc.add_picture(temp_img_path, width=Inches(4))
                        except Exception as e:
                            logging.warning(f"Could not add image: {e}")
                        finally:
                            os.remove(temp_img_path)
                    
                    pix = None
                except Exception as e:
                    logging.warning(f"Could not extract image {img_index}: {e}")
        
        doc.close()
        word_doc.save(output_path)
        return output_path
    
    def pdf_to_excel(self, pdf_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert PDF to Excel spreadsheet (extracts tables and text).
        
        Args:
            pdf_path: Path to input PDF
            output_path: Output path (temp file if None)
            
        Returns:
            Path to created Excel file
        """
        if not self.available_conversions['pdf_to_excel']:
            raise RuntimeError("PDF to Excel conversion not available. Install pdfplumber and openpyxl.")
        
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        if output_path is None:
            output_fd, output_path = tempfile.mkstemp(suffix='.xlsx', prefix='converted_')
            os.close(output_fd)
        
        return self._pdf_to_excel_pdfplumber(pdf_path, output_path)
    
    def _pdf_to_excel_pdfplumber(self, pdf_path: str, output_path: str) -> str:
        """Convert PDF to Excel using pdfplumber."""
        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Content"
        
        # Header style
        header_font = Font(bold=True)
        
        current_row = 1
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Add page header
                ws.cell(row=current_row, column=1, value=f"Page {page_num + 1}").font = header_font
                current_row += 2
                
                # Extract tables first
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables):
                        # Add table header
                        ws.cell(row=current_row, column=1, value=f"Table {table_num + 1}").font = header_font
                        current_row += 1
                        
                        # Add table data
                        for row_data in table:
                            for col_num, cell_data in enumerate(row_data):
                                if cell_data:
                                    ws.cell(row=current_row, column=col_num + 1, value=str(cell_data))
                            current_row += 1
                        current_row += 1  # Space between tables
                
                # Extract remaining text (non-table text)
                text = page.extract_text()
                if text:
                    # Add text header
                    ws.cell(row=current_row, column=1, value="Text Content").font = header_font
                    current_row += 1
                    
                    # Add text content
                    for line in text.split('\n'):
                        if line.strip():
                            ws.cell(row=current_row, column=1, value=line.strip())
                            current_row += 1
                
                current_row += 2  # Space between pages
        
        wb.save(output_path)
        return output_path
    
    def word_to_pdf(self, word_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert Word document to PDF.
        
        Args:
            word_path: Path to input Word document
            output_path: Output path (temp file if None)
            
        Returns:
            Path to created PDF
        """
        if not self.available_conversions['word_to_pdf']:
            raise RuntimeError("Word to PDF conversion not available. Install python-docx and reportlab.")
        
        if not os.path.exists(word_path):
            raise FileNotFoundError(f"Word file not found: {word_path}")
        
        if output_path is None:
            output_fd, output_path = tempfile.mkstemp(suffix='.pdf', prefix='converted_')
            os.close(output_fd)
        
        return self._word_to_pdf_reportlab(word_path, output_path)
    
    def _word_to_pdf_reportlab(self, word_path: str, output_path: str) -> str:
        """Convert Word to PDF using reportlab."""
        doc = Document(word_path)
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine style based on paragraph formatting
                if para.style.name.startswith('Heading'):
                    style = styles['Heading1']
                else:
                    style = styles['Normal']
                
                p = Paragraph(para.text, style)
                story.append(p)
                story.append(Spacer(1, 12))
        
        pdf_doc.build(story)
        return output_path
    
    def excel_to_pdf(self, excel_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert Excel spreadsheet to PDF.
        
        Args:
            excel_path: Path to input Excel file
            output_path: Output path (temp file if None)
            
        Returns:
            Path to created PDF
        """
        if not self.available_conversions['excel_to_pdf']:
            raise RuntimeError("Excel to PDF conversion not available. Install openpyxl and reportlab.")
        
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"Excel file not found: {excel_path}")
        
        if output_path is None:
            output_fd, output_path = tempfile.mkstemp(suffix='.pdf', prefix='converted_')
            os.close(output_fd)
        
        return self._excel_to_pdf_reportlab(excel_path, output_path)
    
    def _excel_to_pdf_reportlab(self, excel_path: str, output_path: str) -> str:
        """Convert Excel to PDF using reportlab."""
        from openpyxl import load_workbook
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors
        
        wb = load_workbook(excel_path)
        pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            # Add sheet title
            styles = getSampleStyleSheet()
            title = Paragraph(f"Sheet: {sheet_name}", styles['Heading1'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Convert sheet data to table
            data = []
            for row in ws.iter_rows(values_only=True):
                # Convert None values to empty strings
                row_data = [str(cell) if cell is not None else '' for cell in row]
                data.append(row_data)
            
            if data:
                # Create table
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                story.append(Spacer(1, 20))
        
        pdf_doc.build(story)
        return output_path
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text content from PDF.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        if PYMUPDF_AVAILABLE:
            return self._extract_text_pymupdf(pdf_path)
        elif PDFPLUMBER_AVAILABLE:
            return self._extract_text_pdfplumber(pdf_path)
        else:
            raise RuntimeError("No text extraction library available")
    
    def _extract_text_pymupdf(self, pdf_path: str) -> str:
        """Extract text using PyMuPDF."""
        doc = fitz.open(pdf_path)
        text = ""
        
        for page in doc:
            text += page.get_text() + "\n\n"
        
        doc.close()
        return text.strip()
    
    def _extract_text_pdfplumber(self, pdf_path: str) -> str:
        """Extract text using pdfplumber."""
        text = ""
        
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        
        return text.strip()
    
    def get_conversion_info(self) -> Dict[str, Any]:
        """
        Get information about available conversions.
        
        Returns:
            Dictionary with conversion capabilities
        """
        return {
            'available_conversions': self.available_conversions,
            'supported_formats': {
                'from_pdf': ['word', 'excel', 'powerpoint', 'text'],
                'to_pdf': ['word', 'excel', 'powerpoint', 'text']
            },
            'requirements': {
                'pdf_to_word': ['PyMuPDF', 'python-docx'],
                'pdf_to_excel': ['pdfplumber', 'openpyxl'],
                'pdf_to_powerpoint': ['PyMuPDF', 'python-pptx'],
                'word_to_pdf': ['python-docx', 'reportlab'],
                'excel_to_pdf': ['openpyxl', 'reportlab'],
                'powerpoint_to_pdf': ['python-pptx', 'reportlab']
            }
        }